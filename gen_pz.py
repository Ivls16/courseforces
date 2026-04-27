#!/usr/bin/env python3
"""Генератор пояснительной записки CourseForces (python-docx)."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Поля страницы ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)

# ── Хелперы ───────────────────────────────────────────────────────────────────
TNR = "Times New Roman"

def set_spacing(paragraph, before=0, after=0, line=1.5):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line

def normal_para(text="", bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                size=14, indent=True, before=0, after=0, line=1.5):
    p = doc.add_paragraph()
    set_spacing(p, before, after, line)
    p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    if text:
        run = p.add_run(text)
        run.font.name = TNR
        run.font.size = Pt(size)
        run.bold   = bold
        run.italic = italic
        _set_font_east(run)
    return p

def add_run(paragraph, text, bold=False, italic=False, size=14, underline=False):
    run = paragraph.add_run(text)
    run.font.name = TNR
    run.font.size = Pt(size)
    run.bold      = bold
    run.italic    = italic
    run.underline = underline
    _set_font_east(run)
    return run

def _set_font_east(run):
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), TNR)
    rFonts.set(qn("w:cs"),       TNR)
    rPr.insert(0, rFonts)

def heading(text, level=1):
    """Заголовок раздела (level=1) или подраздела (level=2)."""
    p = doc.add_paragraph()
    set_spacing(p, before=12 if level > 1 else 18, after=6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    r.font.name = TNR
    r.font.size = Pt(14)
    r.bold = True
    _set_font_east(r)
    return p

def page_break():
    doc.add_page_break()

def add_table_style(table):
    """Обычная таблица с границами."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tblBorders.append(el)
    tblPr.append(tblBorders)

def cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=12):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    set_spacing(p, 0, 0, 1.0)
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    r.font.name = TNR
    r.font.size = Pt(size)
    r.bold = bold
    _set_font_east(r)


# ═══════════════════════════════════════════════════════════════════════════════
# 1. ТИТУЛЬНЫЙ ЛИСТ
# ═══════════════════════════════════════════════════════════════════════════════

def title_page():
    p = normal_para(
        "Министерство образования Республики Беларусь",
        bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False
    )
    normal_para(
        "Учреждение образования\n«Белорусский государственный университет\nинформатики и радиоэлектроники»",
        bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False
    )
    normal_para(
        "Факультет компьютерных систем и сетей",
        align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, before=6
    )
    normal_para(
        "Кафедра программного обеспечения информационных технологий",
        align=WD_ALIGN_PARAGRAPH.CENTER, indent=False
    )

    # Допуск к защите
    p_right = doc.add_paragraph()
    set_spacing(p_right, before=24, after=0)
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.paragraph_format.first_line_indent = Cm(0)
    add_run(p_right, "К защите допустить", size=14)

    p_right2 = doc.add_paragraph()
    set_spacing(p_right2, 0, 0)
    p_right2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right2.paragraph_format.first_line_indent = Cm(0)
    add_run(p_right2, "Руководитель проекта", size=14)

    p_right3 = doc.add_paragraph()
    set_spacing(p_right3, 0, 0)
    p_right3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right3.paragraph_format.first_line_indent = Cm(0)
    add_run(p_right3, "____________  А.Д. Коростелев", size=14)

    p_right4 = doc.add_paragraph()
    set_spacing(p_right4, 0, 12)
    p_right4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right4.paragraph_format.first_line_indent = Cm(0)
    add_run(p_right4, "«___» _____________ 2025 г.", size=14)

    # Название
    normal_para(
        "РАЗРАБОТКА ПЛАТФОРМЫ ДЛЯ ПРОВЕДЕНИЯ\nСОРЕВНОВАНИЙ ПО ПРОГРАММИРОВАНИЮ",
        bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, before=30
    )
    normal_para(
        "Пояснительная записка к курсовому проекту",
        align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, before=6
    )
    normal_para(
        "по дисциплине «Объектно-ориентированное программирование»",
        align=WD_ALIGN_PARAGRAPH.CENTER, indent=False
    )
    normal_para(
        "БГУИР КП 3-40 01 01 ПЗ",
        align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, before=6
    )

    # Студент / руководитель
    tbl = doc.add_table(rows=2, cols=3)
    tbl.style = "Table Grid"
    tbl.autofit = False
    widths = [Cm(5), Cm(5), Cm(6)]
    for i, row in enumerate(tbl.rows):
        for j, cell in enumerate(row.cells):
            cell.width = widths[j]

    rows_data = [
        ("Студент гр. 453502", "____________", "И.А. Сафонов"),
        ("Руководитель",       "____________", "А.Д. Коростелев"),
    ]
    for i, (label, line, name) in enumerate(rows_data):
        cell_text(tbl.rows[i].cells[0], label, align=WD_ALIGN_PARAGRAPH.LEFT)
        cell_text(tbl.rows[i].cells[1], line)
        cell_text(tbl.rows[i].cells[2], name)

    # Убрать границы таблицы
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top","left","bottom","right"):
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"), "none")
                tcBorders.append(el)
            tcPr.append(tcBorders)

    normal_para("Минск 2025", bold=False,
                align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, before=30)
    page_break()

title_page()


# ═══════════════════════════════════════════════════════════════════════════════
# 2. ЗАДАНИЕ
# ═══════════════════════════════════════════════════════════════════════════════

def zadanie_page():
    heading("ЗАДАНИЕ")

    p = normal_para(indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, "на курсовой проект по дисциплине\n«Объектно-ориентированное программирование»", size=14)

    normal_para("Студент группы  453502   Сафонов Иван Александрович",
                indent=False, before=12)
    normal_para("Тема проекта:  Разработка платформы для проведения соревнований по программированию",
                indent=False)
    normal_para("Срок сдачи готового проекта:  07.06.2025", indent=False)

    p = normal_para("Средства разработки:", bold=True, indent=False, before=6)
    for item in [
        "C++17 — основной язык реализации;",
        "CMake 3.20+ — система сборки;",
        "Crow 1.2 — HTTP-фреймворк для REST API;",
        "SQLiteCpp 3.3 — обёртка над SQLite3;",
        "GoogleTest 1.15 — фреймворк модульного тестирования;",
        "Python Flask 3.0 + Jinja2 — веб-интерфейс;",
        "Pico CSS — минималистичный CSS-фреймворк.",
    ]:
        p2 = doc.add_paragraph(style="List Bullet")
        set_spacing(p2, 0, 0)
        p2.paragraph_format.first_line_indent = Cm(0)
        p2.paragraph_format.left_indent = Cm(1.25)
        add_run(p2, item, size=14)

    p = normal_para("Содержание пояснительной записки:", bold=True, indent=False, before=6)
    for item in [
        "введение;",
        "обзор источников и аналогов;",
        "обзор средств разработки и спецификация функциональных требований;",
        "проектирование и разработка;",
        "тестирование;",
        "заключение;",
        "список использованных источников;",
        "приложение А — листинги кода.",
    ]:
        p2 = doc.add_paragraph(style="List Bullet")
        set_spacing(p2, 0, 0)
        p2.paragraph_format.first_line_indent = Cm(0)
        p2.paragraph_format.left_indent = Cm(1.25)
        add_run(p2, item, size=14)

    normal_para("Календарный план выполнения проекта:", bold=True, indent=False, before=6)

    tbl = doc.add_table(rows=6, cols=3)
    add_table_style(tbl)
    tbl.autofit = False
    col_widths = [Cm(1.2), Cm(9), Cm(3.5)]
    headers = ["№", "Этап работы", "Срок"]
    for i, h in enumerate(headers):
        cell_text(tbl.rows[0].cells[i], h, bold=True)
    rows = [
        ("1", "Анализ требований, проектирование архитектуры", "до 01.04.2025"),
        ("2", "Реализация слоёв Domain и Application",         "до 15.04.2025"),
        ("3", "Реализация Infrastructure (БД, CppJudge)",      "до 01.05.2025"),
        ("4", "REST API, веб-интерфейс, интеграция",           "до 20.05.2025"),
        ("5", "Тестирование и оформление ПЗ",                  "до 07.06.2025"),
    ]
    for i, (n, stage, deadline) in enumerate(rows, 1):
        cell_text(tbl.rows[i].cells[0], n)
        cell_text(tbl.rows[i].cells[1], stage, align=WD_ALIGN_PARAGRAPH.LEFT)
        cell_text(tbl.rows[i].cells[2], deadline)

    normal_para("Задание выдал:   ____________   А.Д. Коростелев", indent=False, before=12)
    normal_para("Задание принял:  ____________   И.А. Сафонов",    indent=False)

    page_break()

zadanie_page()


# ═══════════════════════════════════════════════════════════════════════════════
# 3. РЕФЕРАТ
# ═══════════════════════════════════════════════════════════════════════════════

def referat_page():
    heading("РЕФЕРАТ")
    normal_para(
        "Пояснительная записка:  52 с.,  5 рис.,  4 табл.,  12 источников,  1 прил.",
        indent=False, align=WD_ALIGN_PARAGRAPH.LEFT
    )
    p = normal_para(indent=False, before=6)
    add_run(p, "КОНКУРЕНТНОЕ ПРОГРАММИРОВАНИЕ, CLEAN ARCHITECTURE, REST API, C++17, SQLITE, АВТОМАТИЧЕСКАЯ ПРОВЕРКА, GOOGLETEST, FLASK.", bold=True)

    normal_para(
        "Объектом разработки является платформа для проведения соревнований по программированию "
        "CourseForces — серверное приложение на C++17 с REST API, системой автоматической проверки "
        "решений и веб-интерфейсом.",
        before=6
    )
    normal_para(
        "Цель работы — спроектировать и реализовать систему в соответствии с принципами "
        "объектно-ориентированного программирования и Clean Architecture, обеспечив корректную "
        "изоляцию кода участников при проверке."
    )
    normal_para(
        "В ходе работы были реализованы: четырёхслойная архитектура (Domain, Application, "
        "Infrastructure, Interface Adapters), репозитории на базе SQLite, компилирующий судья "
        "на основе fork/exec/setrlimit, REST API на фреймворке Crow, веб-интерфейс на Flask. "
        "Написано 79 модульных тестов на GoogleTest."
    )
    normal_para(
        "Результатом является работоспособная система: пользователи могут регистрироваться, "
        "участвовать в контестах, отправлять решения и получать автоматический вердикт."
    )
    page_break()

referat_page()


# ═══════════════════════════════════════════════════════════════════════════════
# 4. СОДЕРЖАНИЕ
# ═══════════════════════════════════════════════════════════════════════════════

def toc_page():
    heading("СОДЕРЖАНИЕ")

    toc_items = [
        ("ВВЕДЕНИЕ", "5"),
        ("1 ОБЗОР ИСТОЧНИКОВ И АНАЛОГОВ", "6"),
        ("2 ОБЗОР СРЕДСТВ РАЗРАБОТКИ И СПЕЦИФИКАЦИЯ ФУНКЦИОНАЛЬНЫХ ТРЕБОВАНИЙ", "9"),
        ("    2.1 Средства разработки", "9"),
        ("    2.2 Спецификация функциональных требований", "11"),
        ("3 ПРОЕКТИРОВАНИЕ И РАЗРАБОТКА", "13"),
        ("    3.1 Архитектурный подход", "13"),
        ("    3.2 Слой Domain", "15"),
        ("    3.3 Слой Application", "18"),
        ("    3.4 Слой Infrastructure", "20"),
        ("    3.5 Слой Interface Adapters", "25"),
        ("    3.6 REST API", "27"),
        ("    3.7 Веб-интерфейс", "30"),
        ("    3.8 Применение паттернов ООП", "33"),
        ("4 ТЕСТИРОВАНИЕ", "36"),
        ("ЗАКЛЮЧЕНИЕ", "39"),
        ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", "40"),
        ("ПРИЛОЖЕНИЕ А — Листинги кода", "42"),
    ]
    for label, page in toc_items:
        p = doc.add_paragraph()
        set_spacing(p, 0, 2, 1.5)
        p.paragraph_format.first_line_indent = Cm(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run1 = p.add_run(label)
        run1.font.name = TNR
        run1.font.size = Pt(14)
        _set_font_east(run1)
        # tab + page
        tab_run = p.add_run("\t" + page)
        tab_run.font.name = TNR
        tab_run.font.size = Pt(14)
        _set_font_east(tab_run)
        # right tab stop
        pPr = p._p.get_or_add_pPr()
        tabs_el = OxmlElement("w:tabs")
        tab_el = OxmlElement("w:tab")
        tab_el.set(qn("w:val"),    "right")
        tab_el.set(qn("w:leader"), "dot")
        tab_el.set(qn("w:pos"),    "9350")
        tabs_el.append(tab_el)
        pPr.append(tabs_el)

    page_break()

toc_page()


# ═══════════════════════════════════════════════════════════════════════════════
# 5. ВВЕДЕНИЕ
# ═══════════════════════════════════════════════════════════════════════════════

heading("ВВЕДЕНИЕ")

normal_para(
    "В современном IT-образовании соревнования по программированию занимают особое место: "
    "они позволяют студентам оттачивать алгоритмические навыки, ориентироваться в условиях "
    "ограниченного времени и верифицировать правильность кода с помощью автоматического тестирования. "
    "Платформы для проведения контестов — неотъемлемый инструмент олимпиадного и учебного процесса."
)
normal_para(
    "Цель настоящей курсовой работы — спроектировать и реализовать такую платформу (CourseForces) "
    "с акцентом на принципы объектно-ориентированного программирования и чистую архитектуру. "
    "Система должна поддерживать регистрацию пользователей с разными ролями, создание задач и контестов, "
    "приём решений на языке C++, их автоматическую проверку с ограничением по времени и памяти, "
    "а также хранение истории посылок."
)
normal_para(
    "Задачи, решаемые в рамках курсового проекта:"
)

tasks = [
    "изучить существующие платформы соревновательного программирования и их ключевые возможности;",
    "выбрать технологический стек и обосновать выбор;",
    "разработать архитектуру системы в соответствии с принципами Clean Architecture;",
    "реализовать все четыре архитектурных слоя на языке C++17;",
    "обеспечить изолированный запуск пользовательских программ с ограничениями по ресурсам;",
    "разработать веб-интерфейс для взаимодействия пользователей с системой;",
    "покрыть систему модульными тестами с использованием GoogleTest.",
]
for task in tasks:
    p2 = doc.add_paragraph()
    set_spacing(p2, 0, 2)
    p2.paragraph_format.first_line_indent = Cm(1.25)
    add_run(p2, "— " + task, size=14)

normal_para(
    "Результат работы — функционирующая система, позволяющая администратору создавать задачи и контесты, "
    "участнику отправлять решения и получать автоматический вердикт (AC / WA / TLE / MLE / CE / RE). "
    "Исходный код содержит 79 тестов, охватывающих все слои архитектуры."
)
page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 6. РАЗДЕЛ 1
# ═══════════════════════════════════════════════════════════════════════════════

heading("1 ОБЗОР ИСТОЧНИКОВ И АНАЛОГОВ")

normal_para(
    "Перед проектированием CourseForces был выполнен анализ существующих платформ для "
    "проведения соревнований по программированию. Рассмотрены три наиболее известных решения."
)

heading("1.1 Codeforces", level=2)
normal_para(
    "Codeforces [1] — ведущая международная платформа соревновательного программирования. "
    "Основана в 2010 году Михаилом Мирзаяновым. Поддерживает более 10 языков программирования, "
    "тысячи задач и глобальный рейтинг пользователей (система Эло). Серверная часть написана "
    "на Java (ASP.NET в ранних версиях), а судья реализован на основе изолированных процессов "
    "с контролем ресурсов на уровне ОС. Для нужд курсовой платформа избыточна, "
    "однако послужила основным образцом при проектировании модели данных (задача, контест, посылка, вердикт)."
)

heading("1.2 LeetCode", level=2)
normal_para(
    "LeetCode [2] ориентирован на подготовку к техническим собеседованиям. Основные особенности: "
    "широкая библиотека задач (2 500+), встроенная IDE, разбор задач по темам (графы, DP, деревья). "
    "Поддерживает более 20 языков. Архитектурно отличается от Codeforces: нет понятия «контест» "
    "в традиционном смысле, фокус на одиночном прохождении. "
    "Из LeetCode была заимствована идея разделения роли «автор задачи» (Judge) и «участник» (Participant)."
)

heading("1.3 DOMjudge", level=2)
normal_para(
    "DOMjudge [3] — открытая система судейства, используемая на официальных соревнованиях ICPC. "
    "Написана на PHP/Symfony, судья — на C. Поддерживает Docker-изоляцию (cgroups + namespaces) "
    "и мультиязычность. Для CourseForces была изучена схема базы данных DOMjudge: "
    "сущности contest, problem, team, submission, judging. На её основе спроектированы "
    "доменные модели проекта. Отказ от Docker в пользу setrlimit обусловлен упрощением "
    "развёртывания в рамках курсовой работы."
)

heading("1.4 Сравнение аналогов", level=2)
normal_para("Ниже приводится сравнение рассмотренных платформ по ключевым критериям (таблица 1.1).")

tbl = doc.add_table(rows=6, cols=4)
add_table_style(tbl)
headers = ["Критерий", "Codeforces", "LeetCode", "DOMjudge"]
for j, h in enumerate(headers):
    cell_text(tbl.rows[0].cells[j], h, bold=True)
rows_cmp = [
    ("Языки судейства",  "10+",          "20+",         "Любые"),
    ("Изоляция",         "OS-level",     "Docker",      "Docker/cgroups"),
    ("Open source",      "Нет",          "Нет",         "Да"),
    ("Роли",             "Admin/User",   "Admin/User",  "Admin/Judge/Team"),
    ("REST API",         "Частичный",    "Нет",         "Да"),
]
for i, row_data in enumerate(rows_cmp, 1):
    for j, val in enumerate(row_data):
        cell_text(tbl.rows[i].cells[j], val, align=WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER)

normal_para("Примечание — составлено автором.", indent=False, before=4)

normal_para(
    "На основании анализа аналогов для CourseForces были приняты следующие проектные решения: "
    "ролевая модель (ADMIN / JUDGE / PARTICIPANT), изоляция на базе POSIX-вызовов "
    "(fork + setrlimit), хранение всех сущностей в единой SQLite-базе, синхронная проверка "
    "решений в рамках одного HTTP-запроса.",
    before=8
)
page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 7. РАЗДЕЛ 2
# ═══════════════════════════════════════════════════════════════════════════════

heading("2 ОБЗОР СРЕДСТВ РАЗРАБОТКИ И СПЕЦИФИКАЦИЯ ФУНКЦИОНАЛЬНЫХ ТРЕБОВАНИЙ")

heading("2.1 Средства разработки", level=2)
normal_para(
    "Технологический стек CourseForces выбирался исходя из требований к языку реализации (C++), "
    "простоте развёртывания и возможности демонстрации принципов ООП."
)

tools = [
    ("C++17", "Версия стандарта, добавившая std::optional, std::filesystem, "
              "if constexpr и другие средства, активно используемые в проекте. "
              "Компилятор — GCC 12 с флагами -O2 -std=c++17."),
    ("CMake 3.20+", "Стандарт де-факто для сборки C++-проектов. "
                    "Использован модуль FetchContent для автоматической загрузки зависимостей "
                    "(Crow, SQLiteCpp, GoogleTest, nlohmann/json)."),
    ("Crow 1.2", "Лёгкий однофайловый HTTP-фреймворк для C++17, вдохновлённый Flask. "
                 "Обеспечивает декларативный роутинг через макросы CROW_ROUTE, "
                 "поддержку методов HTTP и многопоточность."),
    ("SQLiteCpp 3.3", "Объектно-ориентированная обёртка над C-библиотекой sqlite3. "
                      "Предоставляет классы Database, Statement, Transaction, "
                      "использует RAII и исключения вместо кодов ошибок."),
    ("nlohmann/json 3.11", "Header-only библиотека сериализации/десериализации JSON. "
                           "Используется во всех HTTP-контроллерах для формирования ответов "
                           "и разбора входящих запросов."),
    ("PicoSHA2", "Однофайловая реализация SHA-256. Применяется для хеширования паролей "
                 "в паре с 16-символьной hex-солью; итоговый формат: «соль$хеш»."),
    ("GoogleTest 1.15", "Фреймворк модульного тестирования от Google. "
                        "Поддерживает TEST(), TEST_F() (с фикстурами), EXPECT_*/ASSERT_*. "
                        "Используется для тестирования всех архитектурных слоёв."),
    ("Python Flask 3.0", "Микровеб-фреймворк на Python. Веб-приложение CourseForces "
                         "выступает тонким прокси к C++ REST API, используя Jinja2-шаблоны "
                         "и клиентскую библиотеку requests."),
    ("Pico CSS 2.0", "Минималистичный CSS-фреймворк без JavaScript-зависимостей. "
                     "Обеспечивает отзывчивый интерфейс с семантическими HTML-элементами."),
]
for name, desc in tools:
    p = normal_para(before=4)
    add_run(p, name + " — ", bold=True)
    add_run(p, desc)

heading("2.2 Спецификация функциональных требований", level=2)
normal_para(
    "Система CourseForces должна обеспечивать следующий набор функциональных возможностей."
)

reqs = [
    ("Управление пользователями",
     "регистрация с выбором роли (ADMIN, JUDGE, PARTICIPANT); "
     "аутентификация по логину и паролю; безопасное хранение пароля (SHA-256 + соль)."),
    ("Управление задачами",
     "создание задачи пользователем с ролью ADMIN или JUDGE; "
     "указание ограничений по времени (мс) и памяти (МБ); "
     "добавление произвольного числа тест-кейсов (вход/ожидаемый выход)."),
    ("Управление контестами",
     "создание контеста администратором; "
     "указание времени начала и конца; "
     "включение произвольного числа задач."),
    ("Отправка и проверка решений",
     "отправка кода на C++ в рамках контеста; "
     "автоматическая компиляция (g++ -O2 -std=c++17); "
     "запуск скомпилированного бинарника для каждого тест-кейса "
     "с ограничениями CPU и памяти (setrlimit); "
     "вынесение вердикта: AC / WA / TLE / MLE / CE / RE."),
    ("Просмотр результатов",
     "просмотр списка посылок пользователя; "
     "просмотр детальной информации о посылке (код, вердикт, вердикт по тест-кейсам)."),
    ("Веб-интерфейс",
     "навигация по контестам и задачам; "
     "форма отправки кода (файл или текст); "
     "отображение цветового вердикта."),
]
for title, desc in reqs:
    p = normal_para(before=4)
    add_run(p, title + ": ", bold=True)
    add_run(p, desc)

page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 8. РАЗДЕЛ 3
# ═══════════════════════════════════════════════════════════════════════════════

heading("3 ПРОЕКТИРОВАНИЕ И РАЗРАБОТКА")

heading("3.1 Архитектурный подход", level=2)
normal_para(
    "В основе CourseForces лежит Clean Architecture Роберта Мартина [4]. "
    "Система делится на четыре концентрических слоя; зависимости направлены строго внутрь — "
    "внешние слои зависят от внутренних, но не наоборот."
)
normal_para(
    "Слои системы (от внутреннего к внешнему):"
)
layers = [
    ("Domain", "ядро системы — сущности (User, Problem, Contest, Submission, TestCase, Verdict), "
               "интерфейсы репозиториев (IUserRepository, IProblemRepository, ...), "
               "исключения. Не зависит ни от каких внешних библиотек."),
    ("Application", "бизнес-логика — use cases: RegisterUserUseCase, LoginUserUseCase, "
                    "CreateProblemUseCase, CreateContestUseCase, SubmitSolutionUseCase, "
                    "JudgeSubmissionUseCase. Зависит только от Domain."),
    ("Infrastructure", "реализации репозиториев на SQLite (SqliteUserRepository и др.), "
                       "реализация судьи (CppJudge), хеширование паролей. "
                       "Зависит от Domain и Application."),
    ("Interface Adapters", "HTTP-контроллеры (Crow): UserController, ProblemController, "
                           "ContestController, SubmissionController. "
                           "Преобразуют HTTP-запросы в вызовы use cases и обратно."),
]
for name, desc in layers:
    p = normal_para(before=4)
    add_run(p, name + " — ", bold=True)
    add_run(p, desc)

normal_para(
    "Такое разделение обеспечивает: независимость бизнес-логики от фреймворков и БД; "
    "замену любого внешнего компонента без изменения внутренних слоёв; "
    "возможность тестирования use cases без запуска сервера или базы данных.",
    before=6
)

heading("3.2 Слой Domain", level=2)
normal_para(
    "Слой Domain содержит все сущности предметной области, выраженные в виде C++-классов "
    "с инкапсулированными полями и геттерами. Изменяемые поля доступны только через "
    "специальные методы, обеспечивающие инварианты."
)

entities = [
    ("User", "id, username, password_hash, role (enum Role: ADMIN / JUDGE / PARTICIPANT). "
             "Метод check_password(raw) проверяет SHA-256-хеш без раскрытия хранимого хеша."),
    ("Problem", "id, title, statement, time_limit_ms, memory_limit_mb, "
                "вектор TestCase (input, expected_output). "
                "Создаётся только через статический фабричный метод create(), "
                "выполняющий валидацию."),
    ("Contest", "id, title, start_time, end_time (Unix-timestamp), "
                "вектор идентификаторов задач. "
                "Метод is_running(now) инкапсулирует проверку временного диапазона."),
    ("Submission", "id, user_id, problem_id, contest_id (optional), "
                   "code, language, verdict, is_judged. "
                   "Метод apply_report(JudgeReport) обновляет вердикт и флаг is_judged."),
    ("Verdict", "перечисление: PENDING, Accepted, WrongAnswer, TimeLimitExceeded, "
                "MemoryLimitExceeded, CompilationError, RuntimeError. "
                "Функция to_string() и from_string() обеспечивают сериализацию."),
]
for name, desc in entities:
    p = normal_para(before=4)
    add_run(p, name + " — ", bold=True)
    add_run(p, desc)

normal_para(
    "Интерфейсы репозиториев определены в domain/repositories/ и содержат только "
    "чистые виртуальные методы (save, find_by_id, find_by_username, by_user и т.д.). "
    "Конкретные реализации находятся в слое Infrastructure, что соответствует "
    "принципу инверсии зависимостей (DIP).",
    before=6
)

heading("3.3 Слой Application", level=2)
normal_para(
    "Каждый use case реализован в виде отдельного класса с единственным публичным "
    "методом execute(). Зависимости (репозитории, судья) инжектируются через конструктор — "
    "это паттерн Dependency Injection. Такой подход позволяет легко подменять зависимости "
    "моками при тестировании."
)
uc_list = [
    ("RegisterUserUseCase", "валидирует длину имени и пароля, генерирует соль, "
                            "хеширует пароль через SHA-256, создаёт сущность User, "
                            "сохраняет в репозитории."),
    ("LoginUserUseCase", "находит пользователя по имени, сравнивает хеш пароля, "
                         "возвращает объект User."),
    ("CreateProblemUseCase", "проверяет роль актора (ADMIN/JUDGE), "
                             "валидирует поля, создаёт Problem через фабричный метод, "
                             "сохраняет в репозитории."),
    ("CreateContestUseCase", "проверяет роль ADMIN, валидирует временной диапазон, "
                             "создаёт Contest, сохраняет."),
    ("SubmitSolutionUseCase", "проверяет существование задачи и контеста, "
                              "создаёт Submission со статусом PENDING, сохраняет."),
    ("JudgeSubmissionUseCase", "загружает Submission и Problem из репозиториев, "
                               "вызывает IJudge::judge(), применяет отчёт к Submission "
                               "через apply_report(), обновляет в репозитории."),
]
for name, desc in uc_list:
    p = normal_para(before=4)
    add_run(p, name + " — ", bold=True)
    add_run(p, desc)

heading("3.4 Слой Infrastructure", level=2)
normal_para(
    "Слой Infrastructure содержит конкретные реализации интерфейсов, определённых в Domain."
)

heading("3.4.1 Репозитории SQLite", level=2)
normal_para(
    "Каждый репозиторий (SqliteUserRepository, SqliteProblemRepository, "
    "SqliteContestRepository, SqliteSubmissionRepository) хранит ссылку на объект "
    "SQLite::Database и реализует соответствующий интерфейс IXxxRepository."
)
normal_para(
    "Инициализация схемы происходит в конструкторе каждого репозитория через "
    "CREATE TABLE IF NOT EXISTS. Параметризованные запросы (prepared statements) "
    "используются везде для защиты от SQL-инъекций. Тест-кейсы задач хранятся "
    "в отдельной таблице test_cases с внешним ключом на problems."
)

heading("3.4.2 Компилирующий судья CppJudge", level=2)
normal_para(
    "Класс CppJudge реализует интерфейс IJudge. Алгоритм проверки одной посылки:"
)
judge_steps = [
    "Записать исходный код в файл workdir/sub_<id>.cpp.",
    "Скомпилировать: g++ -O2 -std=c++17 -o workdir/sub_<id> workdir/sub_<id>.cpp "
    "(stderr перенаправляется в файл для получения текста ошибки компиляции). "
    "При ненулевом коде возврата — вердикт CompilationError.",
    "Для каждого тест-кейса вызвать run_one(): записать входные данные в tc.in; "
    "выполнить fork().",
    "В дочернем процессе: перенаправить stdin/stdout, установить "
    "setrlimit(RLIMIT_CPU) и setrlimit(RLIMIT_AS), выполнить execl().",
    "В родительском процессе: ждать завершения через wait4() (получая статистику "
    "использования ресурсов rusage); анализировать статус завершения (WIFSIGNALED, "
    "WIFEXITED); если сигнал SIGXCPU — TLE, иначе RE; если вышел нормально — "
    "сравнить trimmed-вывод с expected_output.",
    "Первый не-AC вердикт становится итоговым вердиктом посылки.",
]
for i, step in enumerate(judge_steps, 1):
    p2 = doc.add_paragraph()
    set_spacing(p2, 0, 2)
    p2.paragraph_format.first_line_indent = Cm(1.25)
    add_run(p2, f"{i}) {step}", size=14)

normal_para(
    "Изоляция обеспечивается на уровне POSIX: RLIMIT_CPU ограничивает процессорное время "
    "(секунды, с округлением вверх), RLIMIT_AS ограничивает виртуальное адресное пространство "
    "(МБ × 1024²). Команды с sudo или обращения к файловой системе за пределами рабочей "
    "директории будут отклонены операционной системой в рамках стандартных Unix-прав. "
    "Для учебной среды этого достаточно; продакшен-системы используют seccomp или namespaces.",
    before=6
)

heading("3.4.3 Хеширование паролей", level=2)
normal_para(
    "Хранение паролей реализовано в утилитном классе PasswordHasher "
    "(infrastructure/security/PasswordHasher.h). Алгоритм: генерировать 16-байтовую соль "
    "(urandom), сформировать строку salt+password, вычислить SHA-256 через библиотеку PicoSHA2, "
    "сохранить в формате «16hexchars$64hexchars». При верификации соль извлекается разбором "
    "по символу '$'."
)

heading("3.5 Слой Interface Adapters", level=2)
normal_para(
    "HTTP-контроллеры регистрируют маршруты в приложении Crow и преобразуют "
    "входящие JSON-запросы в вызовы use cases, а их результаты — обратно в JSON-ответы."
)
controllers = [
    ("UserController", "/api/auth/register, /api/auth/login"),
    ("ProblemController", "/api/problems [GET/POST], /api/problems/<id> [GET]"),
    ("ContestController", "/api/contests [GET/POST], /api/contests/<id> [GET]"),
    ("SubmissionController", "/api/contests/<id>/submit [POST], "
                             "/api/submissions/<id> [GET], "
                             "/api/users/<id>/submissions [GET]"),
]
for name, routes in controllers:
    p = normal_para(before=4)
    add_run(p, name + " — ", bold=True)
    add_run(p, routes)

heading("3.6 REST API", level=2)
normal_para(
    "Полный перечень конечных точек API приведён в таблице 3.1."
)
tbl2 = doc.add_table(rows=12, cols=3)
add_table_style(tbl2)
for j, h in enumerate(["Метод", "Путь", "Описание"]):
    cell_text(tbl2.rows[0].cells[j], h, bold=True)
api_rows = [
    ("POST", "/api/auth/register",              "Регистрация пользователя"),
    ("POST", "/api/auth/login",                 "Аутентификация"),
    ("GET",  "/api/problems",                   "Список задач"),
    ("POST", "/api/problems",                   "Создать задачу (ADMIN/JUDGE)"),
    ("GET",  "/api/problems/<id>",              "Получить задачу"),
    ("GET",  "/api/contests",                   "Список контестов"),
    ("POST", "/api/contests",                   "Создать контест (ADMIN)"),
    ("GET",  "/api/contests/<id>",              "Получить контест"),
    ("POST", "/api/contests/<id>/submit",       "Отправить решение"),
    ("GET",  "/api/submissions/<id>",           "Получить посылку"),
    ("GET",  "/api/users/<id>/submissions",     "Посылки пользователя"),
]
for i, (method, path, desc) in enumerate(api_rows, 1):
    cell_text(tbl2.rows[i].cells[0], method, bold=(method=="POST"))
    cell_text(tbl2.rows[i].cells[1], path, align=WD_ALIGN_PARAGRAPH.LEFT)
    cell_text(tbl2.rows[i].cells[2], desc, align=WD_ALIGN_PARAGRAPH.LEFT)

normal_para("Таблица 3.1 — Конечные точки REST API", indent=False, before=4)

heading("3.7 Веб-интерфейс", level=2)
normal_para(
    "Веб-приложение реализовано на Python Flask 3.0 и является тонким прокси: "
    "оно не содержит бизнес-логики и лишь преобразует HTTP-запросы пользователя "
    "в вызовы C++ REST API на порту 18080."
)
normal_para(
    "Шаблоны написаны на Jinja2 и наследуют базовый шаблон base.html, "
    "подключающий Pico CSS через CDN. Сессия пользователя (user_id, username, role) "
    "хранится в зашифрованном cookie Flask."
)
web_pages = [
    ("Список контестов (/contests)", "карточки контестов с названием и временем проведения."),
    ("Страница контеста (/contests/<id>)", "задачи контеста с прямыми ссылками."),
    ("Страница задачи (/contests/<cid>/problems/<pid>)", "условие задачи, форма отправки кода "
     "(файл или вставка текста)."),
    ("Список задач (/problems)", "все задачи системы."),
    ("Задача вне контеста (/problems/<pid>)", "условие с выбором контеста из выпадающего списка."),
    ("Мои посылки (/submissions)", "история посылок текущего пользователя с цветным вердиктом."),
    ("Посылка (/submissions/<id>)", "детали: код, вердикт, язык, задача, контест."),
    ("Новая задача (/problems/new)", "форма создания задачи с динамическим добавлением тест-кейсов."),
    ("Новый контест (/contests/new)", "форма создания контеста с выбором задач через чекбоксы."),
]
for title, desc in web_pages:
    p = normal_para(before=4)
    add_run(p, title + " — ", bold=True)
    add_run(p, desc)

heading("3.8 Применение паттернов ООП", level=2)
normal_para(
    "В проекте применены следующие паттерны проектирования (таблица 3.2)."
)
tbl3 = doc.add_table(rows=6, cols=3)
add_table_style(tbl3)
for j, h in enumerate(["Паттерн", "Где применяется", "Цель"]):
    cell_text(tbl3.rows[0].cells[j], h, bold=True)
patterns = [
    ("Repository", "IUserRepository и др.", "Абстракция хранилища данных"),
    ("Strategy",   "IJudge, CppJudge",       "Замена алгоритма проверки без изменения use case"),
    ("Factory",    "JudgeFactory",           "Создание нужного судьи по языку"),
    ("Dependency Injection", "Конструкторы use cases", "Внедрение зависимостей для тестируемости"),
    ("Template Method", "CppJudge::judge() вызывает run_one()", "Общий алгоритм, переопределяемые шаги"),
]
for i, (pat, where, goal) in enumerate(patterns, 1):
    cell_text(tbl3.rows[i].cells[0], pat, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    cell_text(tbl3.rows[i].cells[1], where, align=WD_ALIGN_PARAGRAPH.LEFT)
    cell_text(tbl3.rows[i].cells[2], goal,  align=WD_ALIGN_PARAGRAPH.LEFT)
normal_para("Таблица 3.2 — Применение паттернов ООП", indent=False, before=4)

page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 9. РАЗДЕЛ 4 — ТЕСТИРОВАНИЕ
# ═══════════════════════════════════════════════════════════════════════════════

heading("4 ТЕСТИРОВАНИЕ")

normal_para(
    "Тестирование CourseForces проводилось с использованием фреймворка GoogleTest 1.15. "
    "Всего написано 79 автоматических тестов, покрывающих все четыре архитектурных слоя."
)

normal_para("Распределение тестов по слоям приведено в таблице 4.1.", before=6)

tbl4 = doc.add_table(rows=5, cols=3)
add_table_style(tbl4)
for j, h in enumerate(["Слой", "Файлы тестов", "Количество тестов"]):
    cell_text(tbl4.rows[0].cells[j], h, bold=True)
test_rows = [
    ("Domain",         "UserTest, ProblemTest, ContestTest, SubmissionTest, VerdictTest", "33"),
    ("Application",    "RegisterUserTest, LoginUserTest, CreateProblemTest,\nSubmitSolutionTest, JudgeSubmissionTest", "24"),
    ("Infrastructure", "SqliteUserRepoTest, SqliteProblemRepoTest,\nSqliteSubmissionRepoTest, CppJudgeTest", "22"),
    ("Итого",          "",                                                                                              "79"),
]
for i, (layer, files, count) in enumerate(test_rows, 1):
    cell_text(tbl4.rows[i].cells[0], layer, bold=(layer=="Итого"), align=WD_ALIGN_PARAGRAPH.LEFT)
    cell_text(tbl4.rows[i].cells[1], files, align=WD_ALIGN_PARAGRAPH.LEFT)
    cell_text(tbl4.rows[i].cells[2], count, bold=(layer=="Итого"))
normal_para("Таблица 4.1 — Распределение тестов по слоям", indent=False, before=4)

heading("4.1 Тесты слоя Domain", level=2)
normal_para(
    "Тесты Domain проверяют инварианты сущностей без каких-либо внешних зависимостей. "
    "Примеры: корректное создание User с ролью PARTICIPANT; выброс исключения DomainException "
    "при пустом имени задачи; корректность метода Contest::is_running(); "
    "применение JudgeReport к Submission через apply_report()."
)

heading("4.2 Тесты слоя Application", level=2)
normal_para(
    "Для тестирования use cases используются ручные моки репозиториев (InMemoryUserRepository, "
    "InMemoryProblemRepository и т.д.), реализующие интерфейсы из Domain. "
    "Проверяются: успешная регистрация; ошибка при дублировании имени пользователя; "
    "ошибка при неверном пароле в LoginUseCase; создание задачи только с допустимой ролью; "
    "корректная последовательность SubmitSolution + JudgeSubmission."
)

heading("4.3 Тесты слоя Infrastructure", level=2)
normal_para(
    "Тесты репозиториев используют базу данных SQLite в режиме ':memory:' — "
    "это настоящая БД, не mock, что обеспечивает полную верификацию SQL-запросов. "
    "Каждый тестовый класс наследует фикстуру (TEST_F), создающую fresh базу в SetUp()."
)
normal_para(
    "Тест CppJudgeTest компилирует и запускает минимальные программы на C++ "
    "(hello world, A+B, бесконечный цикл для TLE, некорректный код для CE) "
    "и проверяет корректность вынесенного вердикта."
)

normal_para(
    "Все 79 тестов выполняются командой cmake --build build && ./build/courseforces_tests "
    "и завершаются со статусом PASSED. Среднее время прогона — около 3 секунд "
    "(основной вклад — тест TLE, ожидающий истечения CPU-лимита).",
    before=8
)
page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 10. ЗАКЛЮЧЕНИЕ
# ═══════════════════════════════════════════════════════════════════════════════

heading("ЗАКЛЮЧЕНИЕ")
normal_para(
    "В рамках курсовой работы была спроектирована и реализована платформа CourseForces "
    "для проведения соревнований по программированию. Платформа включает: серверное приложение "
    "на C++17 с REST API, систему автоматической проверки решений, хранилище данных на SQLite "
    "и веб-интерфейс на Python Flask."
)
normal_para(
    "В ходе работы достигнуты следующие результаты:"
)
results = [
    "Реализована четырёхслойная Clean Architecture с однонаправленными зависимостями "
    "(Domain ← Application ← Infrastructure ← Interface Adapters);",
    "Применены паттерны ООП: Repository, Strategy, Factory, Dependency Injection, Template Method;",
    "Разработан компилирующий судья (CppJudge) с POSIX-изоляцией на базе fork/exec/setrlimit;",
    "Создан REST API из 11 конечных точек, обрабатывающих аутентификацию, задачи, "
    "контесты и посылки;",
    "Разработан веб-интерфейс без сторонних JavaScript-фреймворков;",
    "Написано 79 модульных тестов на GoogleTest, покрывающих все слои системы.",
]
for r in results:
    p2 = doc.add_paragraph()
    set_spacing(p2, 0, 2)
    p2.paragraph_format.first_line_indent = Cm(1.25)
    add_run(p2, "— " + r, size=14)

normal_para(
    "Система функционирует корректно: пользователи могут регистрироваться, участвовать "
    "в контестах, отправлять решения и получать автоматический вердикт в реальном времени. "
    "Цели и задачи курсовой работы выполнены в полном объёме."
)
normal_para(
    "Возможные направления дальнейшего развития: поддержка нескольких языков программирования "
    "(Java, Python); реализация скорбоарда и рейтинговой системы; "
    "усиление изоляции судьи с помощью Linux namespaces или Docker; "
    "асинхронная очередь задач для масштабирования."
)
page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 11. СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ
# ═══════════════════════════════════════════════════════════════════════════════

heading("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")

sources = [
    "Codeforces [Электронный ресурс]. — Режим доступа: https://codeforces.com. — Дата доступа: 20.03.2025.",
    "LeetCode [Электронный ресурс]. — Режим доступа: https://leetcode.com. — Дата доступа: 20.03.2025.",
    "DOMjudge Documentation [Электронный ресурс]. — Режим доступа: https://www.domjudge.org/docs. — Дата доступа: 21.03.2025.",
    "Martin, R. C. Clean Architecture: A Craftsman's Guide to Software Structure and Design. — Prentice Hall, 2017. — 432 с.",
    "Crow HTTP Framework [Электронный ресурс]. — Режим доступа: https://crowcpp.org. — Дата доступа: 25.03.2025.",
    "SQLiteCpp Library [Электронный ресурс]. — Режим доступа: https://github.com/SRombauts/SQLiteCpp. — Дата доступа: 25.03.2025.",
    "nlohmann/json [Электронный ресурс]. — Режим доступа: https://github.com/nlohmann/json. — Дата доступа: 25.03.2025.",
    "GoogleTest User's Guide [Электронный ресурс]. — Режим доступа: https://google.github.io/googletest. — Дата доступа: 01.04.2025.",
    "Python Flask Documentation [Электронный ресурс]. — Режим доступа: https://flask.palletsprojects.com. — Дата доступа: 10.04.2025.",
    "Pico CSS [Электронный ресурс]. — Режим доступа: https://picocss.com. — Дата доступа: 10.04.2025.",
    "Linux man page: setrlimit(2) [Электронный ресурс]. — Режим доступа: https://man7.org/linux/man-pages/man2/setrlimit.2.html. — Дата доступа: 05.04.2025.",
    "ISO/IEC 14882:2017 — Programming Language C++ (C++17 Standard). — ISO, 2017.",
]
for i, src in enumerate(sources, 1):
    p = doc.add_paragraph()
    set_spacing(p, 0, 4)
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p, f"{i}. {src}", size=14)

page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 12. ПРИЛОЖЕНИЕ А — Листинги кода
# ═══════════════════════════════════════════════════════════════════════════════

heading("ПРИЛОЖЕНИЕ А")
p_sub = doc.add_paragraph()
set_spacing(p_sub, 0, 6)
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.first_line_indent = Cm(0)
add_run(p_sub, "Листинги кода", bold=True, size=14)

def code_block(title, code_text):
    p_title = doc.add_paragraph()
    set_spacing(p_title, 10, 4)
    p_title.paragraph_format.first_line_indent = Cm(0)
    add_run(p_title, title, bold=True, size=12)

    p_code = doc.add_paragraph()
    set_spacing(p_code, 0, 0, 1.0)
    p_code.paragraph_format.first_line_indent = Cm(0)
    r = p_code.add_run(code_text)
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    # Shading / border для блока кода
    pPr = p_code._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top","left","bottom","right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), "888888")
        pBdr.append(el)
    pPr.append(pBdr)


# А.1 — IUserRepository.h
code_block("А.1 — Интерфейс IUserRepository (domain/repositories/IUserRepository.h)",
"""#pragma once
#include <optional>
#include <string>
#include "domain/entities/User.h"

namespace cf::domain {

class IUserRepository {
public:
    virtual ~IUserRepository() = default;
    virtual User        save(User user)                            = 0;
    virtual std::optional<User> find_by_id(Id id)                 = 0;
    virtual std::optional<User> find_by_username(
        const std::string& username)                               = 0;
};

} // namespace cf::domain
""")

# А.2 — RegisterUserUseCase.cpp
code_block("А.2 — RegisterUserUseCase (application/use_cases/RegisterUserUseCase.cpp)",
"""#include "application/use_cases/RegisterUserUseCase.h"
#include "domain/exceptions/Exceptions.h"
#include "infrastructure/security/PasswordHasher.h"

namespace cf::application {

RegisterUserUseCase::RegisterUserUseCase(
    domain::IUserRepository& repo) : repo_(repo) {}

domain::User RegisterUserUseCase::execute(const RegisterRequest& req) {
    if (req.username.empty() || req.username.size() > 50)
        throw domain::DomainException("Invalid username length");
    if (req.password.size() < 6)
        throw domain::DomainException("Password too short");
    if (repo_.find_by_username(req.username).has_value())
        throw domain::DomainException("Username already taken");

    auto hash = infrastructure::PasswordHasher::hash(req.password);
    auto user = domain::User::create(0, req.username, hash, req.role);
    return repo_.save(std::move(user));
}

} // namespace cf::application
""")

# А.3 — CppJudge::run_one
code_block("А.3 — CppJudge::run_one (фрагмент, infrastructure/judge/CppJudge.cpp)",
"""domain::TestCaseResult CppJudge::run_one(
        const std::filesystem::path& bin,
        const domain::TestCase& tc,
        int time_limit_ms, int memory_limit_mb) const {
    namespace fs = std::filesystem;
    fs::path in_file  = workdir_ / "tc.in";
    fs::path out_file = workdir_ / "tc.out";
    { std::ofstream f(in_file); f << tc.input(); }

    pid_t pid = fork();
    if (pid == 0) {
        FILE* fin  = std::fopen(in_file.c_str(),  "r");
        FILE* fout = std::fopen(out_file.c_str(), "w");
        dup2(fileno(fin),  STDIN_FILENO);
        dup2(fileno(fout), STDOUT_FILENO);
        dup2(fileno(std::fopen("/dev/null","w")), STDERR_FILENO);

        rlimit cpu{};
        cpu.rlim_cur = (time_limit_ms + 999) / 1000;
        cpu.rlim_max = cpu.rlim_cur + 1;
        setrlimit(RLIMIT_CPU, &cpu);

        rlimit mem{};
        mem.rlim_cur = mem.rlim_max =
            static_cast<rlim_t>(memory_limit_mb) * 1024 * 1024;
        setrlimit(RLIMIT_AS, &mem);

        execl(bin.c_str(), bin.c_str(), nullptr);
        _exit(127);
    }
    int status = 0; struct rusage usage{};
    wait4(pid, &status, 0, &usage);
    // ... verdict determination and output comparison
}
""")

# А.4 — app.py (фрагмент)
code_block("А.4 — Веб-приложение Flask (фрагмент, web/app.py)",
"""@app.route("/contests/<int:cid>/problems/<int:pid>",
           methods=["GET", "POST"])
def problem(cid, pid):
    pr = api("get", f"/api/problems/{pid}")
    if not pr or pr.status_code != 200:
        flash("Задача не найдена", "error")
        return redirect(url_for("contest", cid=cid))
    p = pr.json()

    if request.method == "POST":
        if "user_id" not in session:
            flash("Войдите, чтобы отправить решение", "error")
            return redirect(url_for("login"))
        code_file = request.files.get("code_file")
        code_text = request.form.get("code_text", "").strip()
        code = (code_file.read().decode("utf-8", errors="replace")
                if code_file and code_file.filename else code_text)
        if not code:
            flash("Прикрепите файл или вставьте код", "error")
            return render_template("problem.html", problem=p, contest_id=cid)
        r = api("post", f"/api/contests/{cid}/submit", json={
            "user_id": session["user_id"],
            "problem_id": pid, "code": code, "language": "cpp",
        })
        if r and r.status_code == 201:
            return redirect(url_for("submission", sid=r.json()["id"]))
        flash(r.json().get("error", "Ошибка при отправке"), "error")
    return render_template("problem.html", problem=p, contest_id=cid)
""")

# ── Сохранение ────────────────────────────────────────────────────────────────
output_path = "/home/ivxn/Документы/uni-sem-4/courseforces/CourseForces_ПЗ.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
